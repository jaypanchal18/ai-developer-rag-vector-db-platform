from fastapi import FastAPI, UploadFile, File, HTTPException
from sqlalchemy import create_engine, Column, Integer, String, LargeBinary
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import os
import PyPDF2
import docx

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://user:password@localhost/dbname")

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, index=True)
    content = Column(LargeBinary)

Base.metadata.create_all(bind=engine)

app = FastAPI()

def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

@app.post("/upload/")
async def upload_document(file: UploadFile = File(...)):
    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF and DOCX are allowed.")
    
    try:
        if file.content_type == "application/pdf":
            text = extract_text_from_pdf(file.file)
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(file.file)
        
        db = SessionLocal()
        db_document = Document(filename=file.filename, content=text.encode('utf-8'))
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        return {"id": db_document.id, "filename": db_document.filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()